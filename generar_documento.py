"""
Genera el documento de presentación del proyecto RehabClinic.
Ejecutar con: python generar_documento.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paleta de colores ─────────────────────────────────────────────────────────
AZUL_OSC   = RGBColor(0x1e, 0x3a, 0x5f)
AZUL_MED   = RGBColor(0x25, 0x63, 0xeb)
AZUL_CLAR  = RGBColor(0xdb, 0xe4, 0xfe)
GRIS_OSC   = RGBColor(0x1e, 0x29, 0x3b)
GRIS_MED   = RGBColor(0x64, 0x74, 0x8b)
GRIS_CLAR  = RGBColor(0xf1, 0xf5, 0xf9)
VERDE      = RGBColor(0x05, 0x96, 0x69)
NARANJA    = RGBColor(0xd9, 0x77, 0x06)
ROJO       = RGBColor(0xdc, 0x26, 0x26)
BLANCO     = RGBColor(0xff, 0xff, 0xff)

OUTPUT = "G:/Otros ordenadores/Gateway/dis+ capacidad/Diseños/Rita Bianchi/listas/RehabClinic - Presentacion del Proyecto.docx"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_hrule(doc, color='2563EB', thickness=12):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    return p

def h1(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSC
    add_hrule(doc, '1e3a5f', 8)
    return p

def h2(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL_MED
    return p

def h3(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_OSC
    return p

def cuerpo(doc, texto, sangria=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    if sangria:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(texto)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_OSC
    return p

def bullet(doc, texto, nivel=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + nivel * 0.6)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(texto)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_OSC
    return p

def nota(doc, texto, color_fondo='EFF6FF', color_borde='2563EB'):
    """Párrafo destacado tipo callout."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, color_fondo)
    set_cell_borders(cell, color_borde)
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(texto)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x5f)
    cell.paragraphs[0].paragraph_format.space_before = Pt(4)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def tabla_simple(doc, filas, anchos_cm, header_bg='1e3a5f'):
    """Tabla con primera fila como encabezado oscuro."""
    n_cols = len(filas[0])
    t = doc.add_table(rows=len(filas), cols=n_cols)
    t.style = 'Table Grid'
    page_w = sum(anchos_cm)

    for r_idx, fila in enumerate(filas):
        row = t.rows[r_idx]
        for c_idx, (texto, bold, color) in enumerate(fila):
            cell = row.cells[c_idx]
            cell.width = Cm(anchos_cm[c_idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                set_cell_bg(cell, header_bg)
                set_cell_borders(cell, header_bg)
            else:
                bg = 'FFFFFF' if r_idx % 2 == 1 else 'F8FAFC'
                set_cell_bg(cell, bg)
                set_cell_borders(cell, 'E2E8F0')
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(texto)
            run.bold = bold
            run.font.size = Pt(10)
            if r_idx == 0:
                run.font.color.rgb = BLANCO
            elif color:
                run.font.color.rgb = color
            else:
                run.font.color.rgb = GRIS_OSC
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.15)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ─────────────────────────────────────────────────────────────────────────────
#  DOCUMENTO
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Márgenes A4
section = doc.sections[0]
section.page_width   = Cm(21)
section.page_height  = Cm(29.7)
section.left_margin  = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin   = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Fuente por defecto
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════

# Bloque azul de portada
t_portada = doc.add_table(rows=1, cols=1)
t_portada.style = 'Table Grid'
c = t_portada.cell(0, 0)
set_cell_bg(c, '1e3a5f')
set_cell_borders(c, '1e3a5f')

for _ in range(2):
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

p_titulo = c.add_paragraph()
p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_titulo.add_run('Sistema de Gestión de Terapias')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = BLANCO

p_sub = c.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_sub.add_run('RehabClinic — Documento de presentación del proyecto')
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

for _ in range(3):
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

doc.add_paragraph()

p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_fecha = p_fecha.add_run('Junio 2026')
r_fecha.font.size = Pt(10)
r_fecha.font.color.rgb = GRIS_MED

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  1. RESUMEN EJECUTIVO
# ════════════════════════════════════════════════════════════════
h1(doc, '1. Resumen ejecutivo')

cuerpo(doc,
    'Este documento describe el Sistema de Gestión de Terapias desarrollado para la institución, '
    'un software diseñado para organizar y optimizar la asignación diaria de sesiones de rehabilitación '
    'entre pacientes y profesionales.')

cuerpo(doc,
    'El sistema reemplaza la planificación manual de la agenda terapéutica por un proceso asistido '
    'por computadora que considera automáticamente las necesidades de cada paciente, la disponibilidad '
    'de los profesionales y las restricciones clínicas definidas por el equipo.')

nota(doc,
    '📌  El objetivo final es que el sistema funcione en la red interna de la institución, '
    'permitiendo que la agenda sea creada desde un equipo centralizado y visualizada en tiempo real '
    'desde cualquier computadora de la clínica, sin necesidad de imprimir ni enviar listas.',
    'EFF6FF', '2563EB')

cuerpo(doc, 'Este documento está organizado en tres partes:')
bullet(doc, 'Lo que el sistema ya hace hoy (versión demostración funcionando).')
bullet(doc, 'Cómo funcionaría el sistema completo en la red de la institución.')
bullet(doc, 'El plan de implementación y los recursos necesarios.')

# ════════════════════════════════════════════════════════════════
#  2. EL SISTEMA HOY — QUÉ HACE LA APLICACIÓN
# ════════════════════════════════════════════════════════════════
h1(doc, '2. El sistema hoy — qué hace la aplicación')

cuerpo(doc,
    'La versión actual es una aplicación web completa que corre en cualquier navegador moderno '
    '(Chrome, Edge, Firefox). No requiere instalación: se abre como si fuera una página de internet, '
    'pero funciona completamente sin conexión a internet.')

h2(doc, '2.1  Módulos disponibles')

tabla_simple(doc, [
    [('Módulo', True, None), ('¿Qué permite hacer?', True, None)],
    [('Agenda del día', True, AZUL_MED),
     ('Ver la grilla completa de sesiones por paciente y horario. Crear, modificar, eliminar e intercambiar sesiones. Generar la agenda automáticamente.', False, None)],
    [('Disponibilidad', True, AZUL_MED),
     ('Marcar qué profesionales están presentes cada día. Registrar horarios bloqueados por paciente y prescripciones urgentes del médico.', False, None)],
    [('Pacientes', True, AZUL_MED),
     ('Ficha completa de cada paciente: grupo diagnóstico, edificio, disciplinas requeridas, bloqueos permanentes de horario, indicadores clínicos.', False, None)],
    [('Profesionales', True, AZUL_MED),
     ('Registro de cada profesional: disciplinas, días laborales, horarios por día, rol de coordinador.', False, None)],
    [('Planes semanales', True, AZUL_MED),
     ('Definir cuántas sesiones de cada disciplina necesita cada paciente por semana. Registro de quién modificó el plan y cuándo.', False, None)],
    [('Métricas', True, AZUL_MED),
     ('Estadísticas de cobertura: horas asignadas vs. requeridas, distribución por disciplina, déficits detectados.', False, None)],
    [('Visualizador', True, AZUL_MED),
     ('Pantalla de solo lectura que muestra la agenda del día con actualización automática. Pensada para televisores o monitores en áreas comunes.', False, None)],
], [5.5, 11.0])

h2(doc, '2.2  El algoritmo de asignación automática')

cuerpo(doc,
    'El corazón del sistema es un algoritmo que resuelve automáticamente el problema de armar la agenda del día. '
    'Este algoritmo trabaja como un rompecabezas inteligente: tiene en cuenta todas las piezas al mismo tiempo '
    'y encuentra la combinación que mejor satisface las necesidades de los pacientes.')

cuerpo(doc, 'Las reglas que considera de forma automática son:')

bullet(doc, 'Prioridad de cada paciente según su nivel de dependencia (asistencia completa, media, mínima, grúa).')
bullet(doc, 'Plan semanal de cada paciente: cuántas sesiones de cada disciplina necesita.')
bullet(doc, 'Horarios bloqueados: tanto los permanentes del perfil del paciente como los coyunturales del día.')
bullet(doc, 'Disponibilidad de los profesionales: solo asigna profesionales que están presentes ese día.')
bullet(doc, 'Horario laboral de cada profesional: días y franjas horarias en que trabaja.')
bullet(doc, 'Preferencia por el profesional de referencia de cada paciente para cada disciplina.')
bullet(doc, 'Orden de terapias: algunas disciplinas rinden más a primera hora, otras a última.')
bullet(doc, 'Kinesiología respiratoria simultánea: permite que un kinesiólogo atienda dos pacientes a la vez si están en el mismo edificio y ninguno tiene la restricción activada.')
bullet(doc, 'Cupo del coordinador: reserva al menos un horario libre por semana para el profesional coordinador.')
bullet(doc, 'Sesiones fijas: respeta sesiones bloqueadas manualmente antes de rearmar el resto.')
bullet(doc, 'Prescripciones urgentes del médico: las asigna con máxima prioridad.')
bullet(doc, 'Rutinas de higiene: reserva automáticamente el primer horario del día para los pacientes que lo requieren.')

nota(doc,
    '💡  Si el resultado automático no es exactamente el deseado, el usuario puede intervenir manualmente: '
    'crear sesiones en celdas vacías, eliminar sesiones, intercambiar horarios entre dos sesiones, '
    'o rotar en cadena tres o más sesiones. El sistema alerta si una modificación viola alguna restricción.',
    'F0FDF4', '059669')

# ════════════════════════════════════════════════════════════════
#  3. EL SISTEMA COMPLETO EN RED
# ════════════════════════════════════════════════════════════════
h1(doc, '3. El sistema completo — funcionamiento en red')

cuerpo(doc,
    'La versión de demostración guarda todos los datos en la memoria del navegador de una sola computadora. '
    'El sistema completo, en cambio, funcionaría en la red interna de la institución, permitiendo que '
    'múltiples equipos accedan a la misma información en tiempo real.')

h2(doc, '3.1  ¿Qué es un servidor y qué hace?')

cuerpo(doc,
    'Un servidor es simplemente una computadora que está siempre encendida y cuya función es '
    '"servir" información a las otras computadoras de la red cuando estas la piden. '
    'No es un equipo especial ni caro: puede ser una PC convencional configurada para este rol, '
    'o uno de los servidores que la institución ya posee.')

cuerpo(doc,
    'En este caso, el servidor tendría dos tareas concretas:')

bullet(doc, 'Guardar y servir los datos del sistema (pacientes, profesionales, agendas, planes).')
bullet(doc, 'Enviar la aplicación al navegador de cualquier computadora de la red cuando alguien la abre.')

nota(doc,
    '🖥️  Analogía: es como el servidor de un restaurante. Las mesas (computadoras) piden cosas, '
    'el mozo (servidor) va a la cocina (disco de datos) y trae lo que se necesita. '
    'Nadie en la mesa tiene que ir a buscar nada; todo pasa por el mozo.',
    'FFFBEB', 'D97706')

h2(doc, '3.2  El disco compartido y los archivos de datos')

cuerpo(doc,
    'Los datos del sistema se guardan como archivos de texto estructurado (formato JSON) '
    'en el disco de acceso general que ya existe en la red de la institución. '
    'Esto tiene varias ventajas importantes:')

bullet(doc, 'El área de IT ya sabe cómo respaldar ese disco: los datos del sistema se incluyen automáticamente en los backups existentes.')
bullet(doc, 'No se necesita una base de datos especial ni licencias adicionales.')
bullet(doc, 'Los archivos son legibles por humanos en caso de emergencia; no son archivos binarios opacos.')
bullet(doc, 'Si el servidor se apaga y se reinicia, los datos persisten intactos en el disco.')

tabla_simple(doc, [
    [('Archivo', True, None), ('Qué contiene', True, None)],
    [('pacientes.json', False, AZUL_MED),     ('Fichas completas de todos los pacientes activos.', False, None)],
    [('profesionales.json', False, AZUL_MED), ('Perfiles de todos los profesionales del equipo.', False, None)],
    [('planes.json', False, AZUL_MED),        ('Planes semanales de sesiones por paciente y disciplina.', False, None)],
    [('asignaciones.json', False, AZUL_MED),  ('Agendas diarias de cada día generado o modificado.', False, None)],
    [('dias_state.json', False, AZUL_MED),    ('Disponibilidad diaria: presencias, bloqueos, prescripciones urgentes.', False, None)],
    [('usuarios.json', False, AZUL_MED),      ('Credenciales de acceso al sistema (contraseñas encriptadas).', False, None)],
], [5.0, 11.5])

h2(doc, '3.3  Las dos aplicaciones')

cuerpo(doc,
    'El sistema completo tiene dos interfaces distintas accesibles desde el navegador:')

h3(doc, 'Aplicación creadora (acceso restringido)')

cuerpo(doc,
    'Es la aplicación completa, con todos los módulos. Solo accesible para usuarios con '
    'credenciales asignadas. Permite crear y modificar toda la información del sistema: '
    'pacientes, profesionales, planes y agendas diarias.')

h3(doc, 'Aplicación visualizadora (acceso libre en la red interna)')

cuerpo(doc,
    'Es una versión de solo lectura que muestra la agenda del día actual. '
    'No requiere usuario ni contraseña: cualquier computadora de la red puede abrirla. '
    'Se actualiza automáticamente cada 30 segundos. Está diseñada para ser '
    'mostrada en pantallas o monitores en áreas comunes del servicio.')

nota(doc,
    '📺  Caso de uso típico: una tablet o monitor en el pasillo del servicio muestra '
    'permanentemente la agenda del día. Cada vez que el coordinador genera o modifica '
    'la agenda desde su equipo, el cambio aparece en todos los monitores en menos de un minuto.',
    'EFF6FF', '2563EB')

h2(doc, '3.4  Diagrama de la arquitectura')

tabla_simple(doc, [
    [('Componente', True, None), ('Ubicación', True, None), ('¿Qué hace?', True, None)],
    [('Servidor HTTP', False, AZUL_MED),
     ('PC servidor o servidor existente', False, None),
     ('Atiende las solicitudes de todas las computadoras de la red y sirve la aplicación.', False, None)],
    [('Disco de datos', False, AZUL_MED),
     ('Disco de acceso general (ya existe)', False, None),
     ('Almacena los archivos JSON con toda la información del sistema.', False, None)],
    [('App creadora', False, AZUL_MED),
     ('Navegador de la PC del coordinador', False, None),
     ('Interfaz completa para crear y gestionar agendas, pacientes y profesionales.', False, None)],
    [('App visualizadora', False, AZUL_MED),
     ('Navegador de cualquier PC / tablet / monitor', False, None),
     ('Muestra la agenda del día en tiempo real. Sin login, solo lectura.', False, None)],
], [4.0, 6.0, 6.5])

# ════════════════════════════════════════════════════════════════
#  4. USUARIOS Y SEGURIDAD
# ════════════════════════════════════════════════════════════════
h1(doc, '4. Usuarios y seguridad')

cuerpo(doc,
    'El sistema tiene un esquema de roles de usuario simple y seguro, '
    'pensado para la realidad operativa de la institución:')

tabla_simple(doc, [
    [('Rol', True, None), ('Permisos', True, None), ('Quiénes lo usan', True, None)],
    [('Administrador', True, VERDE),
     ('Acceso total: crea y gestiona usuarios, modifica todo el sistema.', False, None),
     ('Coordinación o dirección del servicio. Un solo usuario inicial.', False, None)],
    [('Editor', True, AZUL_MED),
     ('Puede crear y modificar agendas, pacientes, profesionales y planes. No puede crear usuarios.', False, None),
     ('Profesionales o administrativos que arman la agenda diaria.', False, None)],
    [('Visualizador', True, GRIS_MED),
     ('Solo lectura. Sin login. Acceso a la app visualizadora desde la red interna.', False, None),
     ('Cualquier PC o monitor en la clínica.', False, None)],
], [3.5, 7.5, 5.5])

cuerpo(doc,
    'Las contraseñas se almacenan de forma encriptada (hash bcrypt) en el archivo usuarios.json. '
    'Ningún usuario puede ver la contraseña de otro, ni siquiera el administrador. '
    'Las sesiones de trabajo expiran automáticamente por inactividad.')

nota(doc,
    '🔒  El sistema solo es accesible desde dentro de la red de la institución. '
    'Ningún dato sale a internet. No hay nube involucrada. Todo queda dentro del edificio.',
    'F0FDF4', '059669')

# ════════════════════════════════════════════════════════════════
#  5. FLUJO DE TRABAJO DIARIO
# ════════════════════════════════════════════════════════════════
h1(doc, '5. Flujo de trabajo diario')

cuerpo(doc, 'Así se usaría el sistema en un día típico de trabajo:')

tabla_simple(doc, [
    [('Paso', True, None), ('Quién', True, None), ('Qué hace', True, None)],
    [('1', False, AZUL_MED), ('Coordinador', False, None),
     ('Abre la app creadora, va a "Disponibilidad" y marca qué profesionales están presentes hoy. Registra ausencias, bloqueos especiales o prescripciones médicas urgentes.', False, None)],
    [('2', False, AZUL_MED), ('Coordinador', False, None),
     ('Hace click en "Generar agenda". El algoritmo crea automáticamente la asignación completa del día en segundos.', False, None)],
    [('3', False, AZUL_MED), ('Coordinador', False, None),
     ('Revisa la agenda generada. Si hay algo que ajustar, lo modifica manualmente: mueve sesiones, crea nuevas, las fija o las elimina.', False, None)],
    [('4', False, VERDE), ('Todos', False, None),
     ('La agenda está visible en tiempo real en la app visualizadora desde cualquier monitor o computadora del servicio.', False, None)],
    [('5', False, GRIS_MED), ('Coordinador', False, None),
     ('A lo largo del día puede hacer ajustes de último momento. Los cambios se reflejan en los monitores en menos de un minuto.', False, None)],
], [1.2, 3.5, 11.8])

# ════════════════════════════════════════════════════════════════
#  6. PLAN DE IMPLEMENTACIÓN
# ════════════════════════════════════════════════════════════════
h1(doc, '6. Plan de implementación')

h2(doc, '6.1  Fases del proyecto')

tabla_simple(doc, [
    [('Fase', True, None), ('Descripción', True, None), ('Estado', True, None)],
    [('Fase 0 — Demostración', True, AZUL_MED),
     ('Aplicación completa funcionando localmente en una sola computadora. Incluye el visualizador de agenda. Esta es la versión presentada en este documento.', False, None),
     ('✅ Completada', True, VERDE)],
    [('Fase 1 — Servidor y red', True, AZUL_MED),
     ('Instalación del servidor HTTP en la red interna. Migración de los datos de localStorage al disco compartido. Acceso desde múltiples PCs simultáneamente.', False, None),
     ('⬜ Pendiente', False, GRIS_MED)],
    [('Fase 2 — Usuarios y login', True, AZUL_MED),
     ('Sistema de autenticación con roles (admin / editor). Registro de quién modificó cada cosa y cuándo.', False, None),
     ('⬜ Pendiente', False, GRIS_MED)],
    [('Fase 3 — Visualizador en red', True, AZUL_MED),
     ('App visualizadora accesible desde cualquier PC o tablet de la institución sin login. Configuración en monitores o televisores del servicio.', False, None),
     ('⬜ Pendiente', False, GRIS_MED)],
    [('Fase 4 — Ajustes y capacitación', True, AZUL_MED),
     ('Ajuste fino del algoritmo según la realidad operativa. Capacitación del equipo. Puesta en producción.', False, None),
     ('⬜ Pendiente', False, GRIS_MED)],
], [4.0, 10.0, 2.5])

h2(doc, '6.2  Requerimientos técnicos')

cuerpo(doc, 'Para implementar el sistema completo en red se necesita:')

h3(doc, 'En el servidor')
bullet(doc, 'Una PC o servidor con Windows que pueda estar encendida permanentemente durante el horario de trabajo.')
bullet(doc, 'Node.js instalado (software gratuito, descarga desde nodejs.org). Ocupa menos de 100 MB.')
bullet(doc, 'Acceso de escritura a una carpeta en el disco de acceso general para guardar los archivos de datos.')
bullet(doc, 'El servidor no necesita ser potente: cualquier PC con 4 GB de RAM y procesador i3 o equivalente es más que suficiente.')

h3(doc, 'En las PCs cliente (las que usan el sistema)')
bullet(doc, 'Solo se necesita un navegador moderno: Chrome, Edge o Firefox. No hay nada que instalar.')
bullet(doc, 'La computadora debe estar conectada a la red de la institución.')
bullet(doc, 'La app se abre escribiendo la dirección del servidor en el navegador (ejemplo: http://192.168.1.10:3000).')

h3(doc, 'Para el visualizador')
bullet(doc, 'Cualquier dispositivo con navegador: PC, tablet, Chromecast con navegador, Smart TV compatible.')
bullet(doc, 'Se accede a la misma red pero a una dirección diferente: http://192.168.1.10:3000/viewer')

nota(doc,
    '💰  Costo de infraestructura adicional: prácticamente nulo. Node.js es gratuito y de código abierto. '
    'No se necesitan licencias de software. Si ya existe un servidor disponible, no se necesita hardware nuevo.',
    'F0FDF4', '059669')

# ════════════════════════════════════════════════════════════════
#  7. BENEFICIOS ESPERADOS
# ════════════════════════════════════════════════════════════════
h1(doc, '7. Beneficios esperados')

tabla_simple(doc, [
    [('Área', True, None), ('Beneficio', True, None)],
    [('Tiempo', False, AZUL_MED),
     ('La generación automática de la agenda diaria pasa de un proceso manual de 30-60 minutos a menos de 10 segundos.', False, None)],
    [('Calidad clínica', False, AZUL_MED),
     ('El algoritmo garantiza que se respeten las necesidades terapéuticas de cada paciente, sin que dependan de la memoria o el criterio subjetivo del día.', False, None)],
    [('Trazabilidad', False, AZUL_MED),
     ('Registro completo de quién hizo qué y cuándo: modificaciones de agenda, cambios en planes, prescripciones médicas.', False, None)],
    [('Visibilidad', False, AZUL_MED),
     ('Todo el equipo ve la agenda actualizada en tiempo real sin depender de impresos, fotocopias ni mensajes de WhatsApp.', False, None)],
    [('Escalabilidad', False, AZUL_MED),
     ('El sistema puede crecer con la institución: más pacientes, más profesionales, nuevas disciplinas, sin cambios estructurales.', False, None)],
    [('Seguridad de datos', False, AZUL_MED),
     ('Los datos se guardan en el disco de la institución, bajo sus propias políticas de backup, sin depender de servicios externos.', False, None)],
], [4.5, 12.0])

# ════════════════════════════════════════════════════════════════
#  8. CIERRE
# ════════════════════════════════════════════════════════════════
h1(doc, '8. Próximos pasos')

cuerpo(doc,
    'La Fase 0 (demostración) está completa y disponible para ser evaluada. '
    'El sistema funciona con datos reales o de ejemplo y puede ser probado por el equipo '
    'antes de tomar cualquier decisión de implementación.')

cuerpo(doc, 'Para avanzar hacia la Fase 1 se necesita:')
bullet(doc, 'Confirmación de la dirección para proceder con el desarrollo.')
bullet(doc, 'Identificación del servidor o PC que actuará como servidor del sistema.')
bullet(doc, 'Coordinación con el área de IT para la configuración de red y permisos de acceso al disco compartido.')
bullet(doc, 'Definición de los usuarios iniciales del sistema y sus roles.')

nota(doc,
    '📋  Todo el código fuente del sistema pertenece a la institución. '
    'No hay dependencia de servicios de terceros, suscripciones ni licencias recurrentes.',
    'EFF6FF', '2563EB')

# ── Guardar ────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f'Documento generado: {OUTPUT}')
